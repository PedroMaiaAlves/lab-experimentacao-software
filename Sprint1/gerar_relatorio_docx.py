"""Gera o relatório intermediário Lab01S03 (RQ01–RQ10) em DOCX."""

from __future__ import annotations

import argparse
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from analysis.constants import LINGUAGENS_POPULARES_OCTOVERSE
from analysis.rq08_popularidadeIntensidade import analisar as analisar_rq08
from analysis.rq09_colaboracao_releases import analisar as analisar_rq09
from analysis.rq10_idade_issues import analisar as analisar_rq10


BASE_DIR = Path(__file__).resolve().parent
CSV_PADRAO = BASE_DIR / "data" / "raw" / "repositorios_populares.csv"
SAIDA_PADRAO = BASE_DIR / "docs" / "Relatorio_Parcial_Lab01S03.docx"
JSON_RQ09_PADRAO = BASE_DIR / "data" / "rq" / "rq09_colaboracao_releases.json"
JSON_RQ10_PADRAO = BASE_DIR / "data" / "rq" / "rq10_idade_issues.json"
TEMPLATE_PADRAO = Path.home() / "Downloads" / "Template_Relatorio_Laboratorio.docx"

AZUL = "17365D"
VERDE = "1F6E63"
VERDE_GRAFICO = "#0F766E"
LARANJA_GRAFICO = "#D97706"
CINZA_GRAFICO = "#6B7280"
CINZA_CLARO = "F0F0EC"
BRANCO = "FFFFFF"
TEXTO = "202020"


def formatar_numero(valor: float | int | None, casas: int = 0) -> str:
    """Formata números no padrão brasileiro sem depender do locale do sistema."""

    if valor is None or pd.isna(valor):
        return "—"
    texto = f"{float(valor):,.{casas}f}"
    return texto.replace(",", "X").replace(".", ",").replace("X", ".")


def formatar_percentual(parte: int | float, total: int | float, casas: int = 1) -> str:
    if not total:
        return "—"
    return f"{formatar_numero(100 * float(parte) / float(total), casas)}%"


def resumo_numerico(serie: pd.Series) -> dict[str, float | int]:
    valores = pd.to_numeric(serie, errors="coerce").dropna()
    return {
        "n": int(len(valores)),
        "media": float(valores.mean()),
        "mediana": float(valores.median()),
        "q1": float(valores.quantile(0.25)),
        "q3": float(valores.quantile(0.75)),
        "minimo": float(valores.min()),
        "maximo": float(valores.max()),
    }


def quantidade_outliers(serie: pd.Series) -> dict[str, float | int]:
    valores = pd.to_numeric(serie, errors="coerce").dropna()
    q1 = float(valores.quantile(0.25))
    q3 = float(valores.quantile(0.75))
    iqr = q3 - q1
    inferior = q1 - 1.5 * iqr
    superior = q3 + 1.5 * iqr
    return {
        "q1": q1,
        "q3": q3,
        "limite_inferior": inferior,
        "limite_superior": superior,
        "inferiores": int((valores < inferior).sum()),
        "superiores": int((valores > superior).sum()),
        "total": int(((valores < inferior) | (valores > superior)).sum()),
    }


def calcular_estatisticas(df: pd.DataFrame) -> dict:
    """Calcula todo número usado no texto, nas tabelas e nas figuras."""

    total = int(len(df))
    idade = pd.to_numeric(df["Idade (anos)"], errors="coerce")
    prs = pd.to_numeric(df["PRs aceitas"], errors="coerce")
    releases = pd.to_numeric(df["Total de releases"], errors="coerce")
    dias = pd.to_numeric(df["Dias desde última atualização"], errors="coerce")
    issues = pd.to_numeric(df["% issues fechadas"], errors="coerce")
    linguagens = df["Linguagem"].fillna("Não definida")
    top10 = set(LINGUAGENS_POPULARES_OCTOVERSE)

    faixas_idade = [
        int((idade < 2).sum()),
        int(((idade >= 2) & (idade < 5)).sum()),
        int(((idade >= 5) & (idade < 10)).sum()),
        int((idade >= 10).sum()),
    ]
    faixas_dias = [
        int((dias <= 1).sum()),
        int(((dias > 1) & (dias <= 7)).sum()),
        int(((dias > 7) & (dias <= 30)).sum()),
        int(((dias > 30) & (dias <= 365)).sum()),
        int((dias > 365).sum()),
    ]

    contagem_linguagens = linguagens.value_counts()
    quantidade_top10 = int(linguagens.isin(top10).sum())
    quantidade_indefinida = int((linguagens == "Não definida").sum())
    quantidade_definida = total - quantidade_indefinida
    outras_definidas = quantidade_definida - quantidade_top10

    definidos = df[linguagens != "Não definida"].copy()
    definidos["Grupo"] = np.where(
        definidos["Linguagem"].isin(top10), "Top 10 Octoverse", "Outras definidas"
    )
    rq07 = (
        definidos.groupby("Grupo")
        .agg(
            n=("Nome", "size"),
            prs=("PRs aceitas", "median"),
            releases=("Total de releases", "median"),
            dias=("Dias desde última atualização", "median"),
        )
        .reindex(["Top 10 Octoverse", "Outras definidas"])
    )

    rq08 = analisar_rq08(df)
    rq09 = analisar_rq09(df)
    rq10 = analisar_rq10(df)
    total_faixas_rq09 = sum(
        faixa["repositorios"] for faixa in rq09["resumo_por_quartil_prs"]
    )
    if rq09["total_repositorios"] != total:
        raise RuntimeError("A RQ09 não cobre o mesmo total de repositórios do CSV.")
    if total_faixas_rq09 != rq09["amostra"]["analise_principal"]:
        raise RuntimeError("Os quartis da RQ09 não cobrem a amostra principal.")
    if rq09["correlacao_spearman"]["n"] != rq09["amostra"]["analise_principal"]:
        raise RuntimeError("A correlação da RQ09 não usa toda a amostra principal.")
    rho_rq09 = rq09["correlacao_spearman"]["rho"]
    if rho_rq09 is not None and not -1 <= rho_rq09 <= 1:
        raise RuntimeError("A correlação da RQ09 está fora do intervalo [-1, 1].")
    total_faixas_rq10 = sum(
        faixa["Quantidade_repositorios"]
        for faixa in rq10["resumo_por_faixa_etaria"]
    )
    if rq10["total_repositorios"] != total:
        raise RuntimeError("A RQ10 não cobre o mesmo total de repositórios do CSV.")
    if rq10["repositorios_com_issues"] + rq10["repositorios_sem_issues"] != total:
        raise RuntimeError("As amostras com e sem issues da RQ10 não somam o total.")
    if total_faixas_rq10 != rq10["repositorios_com_issues"]:
        raise RuntimeError("As faixas etárias da RQ10 não cobrem a amostra válida.")
    rho_rq10 = rq10["correlacao_spearman"]
    if rho_rq10 is not None and not -1 <= rho_rq10 <= 1:
        raise RuntimeError("A correlação da RQ10 está fora do intervalo [-1, 1].")
    return {
        "total": total,
        "idade": resumo_numerico(idade),
        "faixas_idade": faixas_idade,
        "prs": resumo_numerico(prs),
        "prs_zero": int((prs == 0).sum()),
        "prs_outliers": quantidade_outliers(prs),
        "top_prs": df.nlargest(3, "PRs aceitas")[["Nome", "PRs aceitas"]].values.tolist(),
        "releases": resumo_numerico(releases),
        "releases_zero": int((releases == 0).sum()),
        "releases_com": int((releases > 0).sum()),
        "releases_outliers": quantidade_outliers(releases),
        "top_releases": df.nlargest(3, "Total de releases")[["Nome", "Total de releases"]].values.tolist(),
        "dias": resumo_numerico(dias),
        "faixas_dias": faixas_dias,
        "dias_ate_30": int((dias <= 30).sum()),
        "dias_outliers": quantidade_outliers(dias),
        "linguagens_distintas_definidas": int(linguagens[linguagens != "Não definida"].nunique()),
        "linguagens_top": list(contagem_linguagens.head(8).items()),
        "top10": quantidade_top10,
        "outras_definidas": outras_definidas,
        "indefinida": quantidade_indefinida,
        "definida": quantidade_definida,
        "issues": resumo_numerico(issues),
        "issues_sem": int(issues.isna().sum()),
        "issues_ge80": int((issues >= 80).sum()),
        "issues_lt50": int((issues < 50).sum()),
        "issues_100": int((issues == 100).sum()),
        "issues_outliers": quantidade_outliers(issues),
        "rq07": rq07,
        "rq08": rq08,
        "rq09": rq09,
        "rq10": rq10,
    }


def limpar_corpo(documento: Document) -> None:
    body = documento._element.body
    for filho in list(body):
        if filho.tag != qn("w:sectPr"):
            body.remove(filho)


def configurar_estilos(documento: Document) -> None:
    normal = documento.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXTO)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for nome, tamanho, cor in (
        ("Title", 26, AZUL),
        ("Subtitle", 12, VERDE),
        ("Heading 1", 16, AZUL),
        ("Heading 2", 13, VERDE),
        ("Heading 3", 11.5, AZUL),
    ):
        estilo = documento.styles[nome]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(tamanho)
        estilo.font.bold = nome != "Subtitle"
        estilo.font.color.rgb = RGBColor.from_string(cor)
        estilo.paragraph_format.keep_with_next = True


def sombrear(celula, cor: str) -> None:
    tc_pr = celula._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), cor)


def adicionar_numero_pagina(paragrafo) -> None:
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run("Página ")
    for tipo, texto in (("begin", None), (None, " PAGE "), ("end", None)):
        elemento = OxmlElement("w:instrText" if texto else "w:fldChar")
        if texto:
            elemento.set(qn("xml:space"), "preserve")
            elemento.text = texto
        else:
            elemento.set(qn("w:fldCharType"), tipo)
        run._r.append(elemento)


def adicionar_texto(documento: Document, texto: str, negrito_inicial: str | None = None):
    paragrafo = documento.add_paragraph(style="Normal")
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if negrito_inicial and texto.startswith(negrito_inicial):
        paragrafo.add_run(negrito_inicial).bold = True
        paragrafo.add_run(texto[len(negrito_inicial) :])
    else:
        paragrafo.add_run(texto)
    return paragrafo


def adicionar_lista(documento: Document, itens: list[str]) -> None:
    for item in itens:
        documento.add_paragraph(item, style="List Bullet")


def adicionar_tabela(documento: Document, cabecalhos: list[str], linhas: list[list[object]], fonte: float = 9):
    tabela = documento.add_table(rows=1, cols=len(cabecalhos))
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    for indice, cabecalho in enumerate(cabecalhos):
        celula = tabela.rows[0].cells[indice]
        sombrear(celula, VERDE)
        celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = celula.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(cabecalho))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(BRANCO)
        run.font.size = Pt(fonte)
    tr_pr = tabela.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    for linha in linhas:
        celulas = tabela.add_row().cells
        for indice, valor in enumerate(linha):
            celulas[indice].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = celulas[indice].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if indice else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(str(valor)).font.size = Pt(fonte)
    documento.add_paragraph()
    return tabela


def adicionar_figura(documento: Document, caminho: Path, numero: int, legenda: str) -> None:
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(caminho), width=Cm(15.8))
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figura {numero} — {legenda}")
    run.bold = True
    run.font.size = Pt(9.5)
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fonte: elaborado pelo grupo com dados da API GraphQL do GitHub, 2026.")
    run.italic = True
    run.font.size = Pt(8.5)


def salvar_figura(fig, caminho: Path) -> Path:
    fig.tight_layout()
    fig.savefig(caminho, dpi=220, bbox_inches="tight")
    plt.close(fig)
    return caminho


def preparar_graficos(df: pd.DataFrame, estatisticas: dict, diretorio: Path) -> dict[str, Path]:
    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 9, "axes.grid": False})
    graficos: dict[str, Path] = {}
    total = estatisticas["total"]

    fig, ax = plt.subplots(figsize=(8.2, 4.5))
    ax.hist(df["Idade (anos)"], bins=np.arange(0, max(22, math.ceil(df["Idade (anos)"].max()) + 2), 2), color=VERDE_GRAFICO, edgecolor="white")
    mediana = estatisticas["idade"]["mediana"]
    ax.axvline(mediana, color=LARANJA_GRAFICO, linestyle="--", linewidth=2, label=f"Mediana: {formatar_numero(mediana, 2)} anos")
    ax.set(title="Distribuição da idade dos repositórios", xlabel="Idade (anos)", ylabel="Repositórios")
    ax.legend(frameon=False)
    ax.text(0.99, 0.95, f"n = {formatar_numero(total)}", transform=ax.transAxes, ha="right")
    graficos["rq01"] = salvar_figura(fig, diretorio / "rq01.png")

    prs = pd.to_numeric(df["PRs aceitas"], errors="coerce")
    limites_pr = [0, 1, 10, 100, 1_000, 10_000, 100_000, np.inf]
    rotulos_pr = ["0", "1–9", "10–99", "100–999", "1.000–9.999", "10.000–99.999", "≥100.000"]
    valores_pr = [int(((prs >= a) & (prs < b)).sum()) for a, b in zip(limites_pr[:-1], limites_pr[1:])]
    fig, ax = plt.subplots(figsize=(8.6, 4.6))
    barras = ax.bar(rotulos_pr, valores_pr, color=VERDE_GRAFICO)
    ax.bar_label(barras, padding=3)
    ax.set(title="Repositórios por faixa de PRs mescladas", xlabel="PRs mescladas acumuladas", ylabel="Repositórios")
    ax.tick_params(axis="x", rotation=20)
    ax.text(0.99, 0.94, f"Mediana: {formatar_numero(estatisticas['prs']['mediana'])} | Máximo: {formatar_numero(estatisticas['prs']['maximo'])}", transform=ax.transAxes, ha="right", fontweight="bold")
    graficos["rq02"] = salvar_figura(fig, diretorio / "rq02.png")

    releases = pd.to_numeric(df["Total de releases"], errors="coerce")
    limites_rel = [0, 1, 10, 50, 100, 500, 1_000, np.inf]
    rotulos_rel = ["0", "1–9", "10–49", "50–99", "100–499", "500–999", "≥1.000"]
    valores_rel = [int(((releases >= a) & (releases < b)).sum()) for a, b in zip(limites_rel[:-1], limites_rel[1:])]
    fig, ax = plt.subplots(figsize=(8.4, 4.6))
    barras = ax.bar(rotulos_rel, valores_rel, color=VERDE_GRAFICO)
    ax.bar_label(barras, padding=3)
    ax.set(title="Repositórios por quantidade acumulada de releases", xlabel="Releases acumuladas", ylabel="Repositórios")
    ax.text(0.99, 0.94, f"Mediana: {formatar_numero(estatisticas['releases']['mediana'])} | Com release: {formatar_numero(estatisticas['releases_com'])}", transform=ax.transAxes, ha="right", fontweight="bold")
    graficos["rq03"] = salvar_figura(fig, diretorio / "rq03.png")

    rotulos_dias = ["Até 1 dia", ">1–7", ">7–30", ">30–365", ">365"]
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    barras = ax.bar(rotulos_dias, estatisticas["faixas_dias"], color=VERDE_GRAFICO)
    ax.bar_label(barras, labels=[f"{v}\n({formatar_percentual(v, total)})" for v in estatisticas["faixas_dias"]], padding=3)
    ax.set(title="Tempo desde o último push", xlabel="Faixa de recência", ylabel="Repositórios")
    graficos["rq04"] = salvar_figura(fig, diretorio / "rq04.png")

    valores_ling = [estatisticas["top10"], estatisticas["outras_definidas"], estatisticas["indefinida"]]
    nomes_ling = ["Top 10 Octoverse", "Outras definidas", "Não definida"]
    fig, ax = plt.subplots(figsize=(8.5, 3.1))
    esquerda = 0
    for valor, nome, cor in zip(valores_ling, nomes_ling, [VERDE_GRAFICO, LARANJA_GRAFICO, CINZA_GRAFICO]):
        ax.barh([f"{formatar_numero(total)} repositórios"], [valor], left=esquerda, color=cor, label=nome)
        if valor:
            ax.text(esquerda + valor / 2, 0, f"{valor}\n{formatar_percentual(valor, total)}", ha="center", va="center", color="white", fontweight="bold")
        esquerda += valor
    ax.set_xlim(0, total)
    ax.set_title("Presença das linguagens do Top 10 Octoverse")
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.24), ncol=3, frameon=False)
    graficos["rq05"] = salvar_figura(fig, diretorio / "rq05.png")

    issues = pd.to_numeric(df["% issues fechadas"], errors="coerce")
    validos = issues.dropna()
    faixas_issues = [int(((validos >= x) & (validos < x + 10)).sum()) for x in range(0, 90, 10)]
    faixas_issues += [int(((validos >= 90) & (validos <= 100)).sum()), int(issues.isna().sum())]
    fig, ax = plt.subplots(figsize=(9, 4.6))
    barras = ax.bar([f"{x}–<{x + 10}%" for x in range(0, 90, 10)] + ["90–100%", "Sem issues"], faixas_issues, color=[VERDE_GRAFICO] * 10 + [CINZA_GRAFICO])
    ax.bar_label(barras, padding=3, fontsize=8)
    ax.tick_params(axis="x", rotation=32)
    ax.set(title="Distribuição do percentual de issues fechadas", xlabel="Percentual", ylabel="Repositórios")
    graficos["rq06"] = salvar_figura(fig, diretorio / "rq06.png")

    rq07 = estatisticas["rq07"]
    fig, eixos = plt.subplots(1, 3, figsize=(10.2, 4.2))
    for ax, (coluna, titulo) in zip(eixos, [("prs", "Mediana de PRs"), ("releases", "Mediana de releases"), ("dias", "Mediana de dias")]):
        valores = rq07[coluna].astype(float).values
        barras = ax.bar(["Top 10", "Outras"], valores, color=[VERDE_GRAFICO, LARANJA_GRAFICO])
        ax.bar_label(barras, labels=[formatar_numero(v, 2 if coluna == "dias" else 0) for v in valores], padding=3)
        ax.set_title(titulo)
    fig.suptitle("Indicadores medianos por grupo de linguagens", fontweight="bold")
    graficos["rq07"] = salvar_figura(fig, diretorio / "rq07.png")

    principal = df[pd.to_numeric(df["Idade (anos)"], errors="coerce") >= 1].copy()
    rq08 = estatisticas["rq08"]
    correlacoes = rq08["correlacoes_spearman"]["principal_idade_maior_igual_1"]
    fig, eixos = plt.subplots(1, 2, figsize=(10.4, 4.6))
    for ax, coluna, titulo, chave in (
        (eixos[0], "PRs por ano", "Estrelas × PRs mescladas/ano", "estrelas_vs_prs_por_ano"),
        (eixos[1], "Releases por ano", "Estrelas × releases/ano", "estrelas_vs_releases_por_ano"),
    ):
        ax.scatter(principal["Estrelas"], principal[coluna], s=14, alpha=0.25, color=VERDE_GRAFICO, edgecolors="none")
        ax.set_xscale("log")
        ax.set_yscale("symlog", linthresh=1)
        ax.set(title=titulo, xlabel="Estrelas (log)", ylabel=f"{coluna} (symlog)")
        resultado = correlacoes[chave]
        ax.text(0.04, 0.95, f"ρ = {formatar_numero(resultado['rho'], 4)}\nn = {resultado['n']}", transform=ax.transAxes, va="top", bbox={"facecolor": "white", "alpha": 0.85, "edgecolor": "none"})
    fig.suptitle("RQ08 — Popularidade e intensidade anual (idade ≥ 1 ano)", fontweight="bold")
    graficos["rq08"] = salvar_figura(fig, diretorio / "rq08.png")

    rq09 = estatisticas["rq09"]
    limites_prs = rq09["outliers_iqr"]["prs_por_ano"]
    limites_releases = rq09["outliers_iqr"]["releases_por_ano"]
    mascara_outliers = (
        (principal["PRs por ano"] < limites_prs["limite_inferior"])
        | (principal["PRs por ano"] > limites_prs["limite_superior"])
        | (principal["Releases por ano"] < limites_releases["limite_inferior"])
        | (principal["Releases por ano"] > limites_releases["limite_superior"])
    )
    fig, ax = plt.subplots(figsize=(9.2, 5.2))
    ax.scatter(
        principal.loc[~mascara_outliers, "PRs por ano"],
        principal.loc[~mascara_outliers, "Releases por ano"],
        s=16,
        alpha=0.28,
        color=VERDE_GRAFICO,
        edgecolors="none",
        label="Demais repositórios",
    )
    ax.scatter(
        principal.loc[mascara_outliers, "PRs por ano"],
        principal.loc[mascara_outliers, "Releases por ano"],
        s=22,
        alpha=0.7,
        color=LARANJA_GRAFICO,
        edgecolors="none",
        label="Outlier em ao menos uma taxa (IQR)",
    )
    coeficiente, intercepto = np.polyfit(
        principal["PRs por ano"], principal["Releases por ano"], 1
    )
    eixo_tendencia = np.linspace(
        principal["PRs por ano"].min(), principal["PRs por ano"].max(), 200
    )
    ax.plot(
        eixo_tendencia,
        coeficiente * eixo_tendencia + intercepto,
        color="#B91C1C",
        linewidth=2,
        label="Tendência linear",
    )
    ax.set_xscale("symlog", linthresh=1)
    ax.set_yscale("symlog", linthresh=1)
    ax.set_ylim(bottom=0)
    ax.set(
        title="Intensidade de PRs mescladas × intensidade de releases",
        xlabel="PRs mescladas por ano (symlog)",
        ylabel="Releases por ano (symlog)",
    )
    correlacao_rq09 = rq09["correlacao_spearman"]
    ax.text(
        0.04,
        0.95,
        f"ρ = {formatar_numero(correlacao_rq09['rho'], 4)}\n"
        f"n = {formatar_numero(correlacao_rq09['n'])}",
        transform=ax.transAxes,
        va="top",
        bbox={"facecolor": "white", "alpha": 0.85, "edgecolor": "none"},
    )
    ax.legend(frameon=False, loc="lower right")
    fig.suptitle("RQ09 — Colaboração e publicação de releases", fontweight="bold")
    graficos["rq09"] = salvar_figura(fig, diretorio / "rq09.png")

    rq10 = estatisticas["rq10"]
    dados_rq10 = df[["Idade (anos)", "% issues fechadas"]].copy()
    dados_rq10["Idade (anos)"] = pd.to_numeric(
        dados_rq10["Idade (anos)"], errors="coerce"
    )
    dados_rq10["% issues fechadas"] = pd.to_numeric(
        dados_rq10["% issues fechadas"], errors="coerce"
    )
    dados_rq10 = dados_rq10.dropna()
    resumo_rq10 = pd.DataFrame(rq10["resumo_por_faixa_etaria"])

    fig, eixos = plt.subplots(1, 2, figsize=(10.6, 4.8))
    eixos[0].scatter(
        dados_rq10["Idade (anos)"],
        dados_rq10["% issues fechadas"],
        s=15,
        alpha=0.22,
        color=VERDE_GRAFICO,
        edgecolors="none",
    )
    eixos[0].set(
        title="Idade × percentual de issues fechadas",
        xlabel="Idade (anos)",
        ylabel="Issues fechadas (%)",
        ylim=(-2, 102),
    )
    eixos[0].text(
        0.04,
        0.95,
        f"ρ = {formatar_numero(rq10['correlacao_spearman'], 4)}\n"
        f"n = {formatar_numero(rq10['repositorios_com_issues'])}",
        transform=eixos[0].transAxes,
        va="top",
        bbox={"facecolor": "white", "alpha": 0.85, "edgecolor": "none"},
    )

    posicoes = np.arange(len(resumo_rq10))
    medianas = resumo_rq10["Mediana_issues_fechadas"].to_numpy(dtype=float)
    q1 = resumo_rq10["Q1_issues_fechadas"].to_numpy(dtype=float)
    q3 = resumo_rq10["Q3_issues_fechadas"].to_numpy(dtype=float)
    eixos[1].errorbar(
        posicoes,
        medianas,
        yerr=np.vstack((medianas - q1, q3 - medianas)),
        fmt="o-",
        color=VERDE_GRAFICO,
        ecolor=LARANJA_GRAFICO,
        capsize=5,
        linewidth=2,
        markersize=6,
    )
    for posicao, mediana_faixa, quantidade in zip(
        posicoes,
        medianas,
        resumo_rq10["Quantidade_repositorios"],
    ):
        eixos[1].annotate(
            f"{formatar_numero(mediana_faixa, 2)}%\nn={int(quantidade)}",
            (posicao, mediana_faixa),
            xytext=(0, 9),
            textcoords="offset points",
            ha="center",
            fontsize=8,
        )
    eixos[1].set(
        title="Mediana e intervalo interquartil por idade",
        xlabel="Faixa de idade",
        ylabel="Issues fechadas (%)",
        ylim=(45, 105),
        xticks=posicoes,
        xticklabels=resumo_rq10["Faixa de idade"],
    )
    eixos[1].tick_params(axis="x", rotation=18)
    fig.suptitle("RQ10 — Maturidade e tratamento de issues", fontweight="bold")
    graficos["rq10"] = salvar_figura(fig, diretorio / "rq10.png")
    return graficos


def validar_dados(df: pd.DataFrame) -> None:
    colunas = {
        "Nome", "URL", "Estrelas", "Linguagem", "Idade (anos)", "PRs aceitas",
        "Total de releases", "PRs por ano", "Releases por ano",
        "Dias desde última atualização", "Issues abertas", "Issues fechadas",
        "% issues fechadas",
    }
    ausentes = colunas.difference(df.columns)
    if ausentes:
        raise RuntimeError(f"CSV sem as colunas esperadas: {sorted(ausentes)}")
    if len(df) != 1_000:
        raise RuntimeError(f"O relatório exige 1.000 registros; o CSV contém {len(df)}.")
    if df["Nome"].nunique() != 1_000 or df["URL"].nunique() != 1_000:
        raise RuntimeError("O CSV não contém 1.000 nomes e URLs únicos.")
    if not pd.to_numeric(df["Estrelas"], errors="coerce").is_monotonic_decreasing:
        raise RuntimeError("O CSV não está ordenado por estrelas de forma não crescente.")
    for coluna in [
        "Estrelas",
        "Idade (anos)",
        "PRs aceitas",
        "Total de releases",
        "PRs por ano",
        "Releases por ano",
        "Issues abertas",
        "Issues fechadas",
    ]:
        valores = pd.to_numeric(df[coluna], errors="coerce")
        validos = valores.dropna()
        if (validos < 0).any() or not np.isfinite(validos).all():
            raise RuntimeError(f"A coluna {coluna} contém valor negativo ou não finito.")
    percentual_issues = pd.to_numeric(df["% issues fechadas"], errors="coerce")
    percentual_valido = percentual_issues.dropna()
    if (
        not np.isfinite(percentual_valido).all()
        or (percentual_valido < 0).any()
        or (percentual_valido > 100).any()
    ):
        raise RuntimeError("O percentual de issues fechadas está fora de 0–100.")
    abertas = pd.to_numeric(df["Issues abertas"], errors="coerce")
    fechadas = pd.to_numeric(df["Issues fechadas"], errors="coerce")
    total_issues = abertas + fechadas
    if percentual_issues[total_issues == 0].notna().any():
        raise RuntimeError("Projetos sem issues devem ter percentual de fechamento nulo.")
    if percentual_issues[total_issues > 0].isna().any():
        raise RuntimeError("Projetos com issues devem ter percentual de fechamento definido.")
    idade = pd.to_numeric(df["Idade (anos)"], errors="coerce")
    esperado_prs = (pd.to_numeric(df["PRs aceitas"], errors="coerce") / idade).round(4).where(idade > 0)
    esperado_releases = (pd.to_numeric(df["Total de releases"], errors="coerce") / idade).round(4).where(idade > 0)
    if not np.allclose(df["PRs por ano"].fillna(-1), esperado_prs.fillna(-1), atol=1e-4):
        raise RuntimeError("PRs por ano não corresponde às contagens divididas pela idade registrada.")
    if not np.allclose(df["Releases por ano"].fillna(-1), esperado_releases.fillna(-1), atol=1e-4):
        raise RuntimeError("Releases por ano não corresponde às contagens divididas pela idade registrada.")


def adicionar_capa(documento: Document) -> None:
    p = documento.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO DE LABORATÓRIO")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(AZUL)
    p = documento.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mineração e análise dos 1.000 repositórios mais populares do GitHub")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(VERDE)
    p = documento.add_paragraph("Lab01S03 — versão intermediária com RQ01–RQ10")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    dados = [
        ["Curso", "Engenharia de Software"],
        ["Disciplina", "Laboratório de Experimentação de Software"],
        ["Professor", "Danilo Maia"],
        ["Entrega", "Lab01S03 — análise e visualização de dados"],
        ["Integrantes", "Pedro Henrique Maia Alves; Diogo C. Brunoro; Lorran Pedro Avelar Xavier"],
        ["Situação", "RQ01–RQ10 integradas; snapshot final do board pendente"],
        ["Data de geração", date.today().strftime("%d/%m/%Y")],
    ]
    tabela = adicionar_tabela(documento, ["Identificação", "Informação"], dados, fonte=9.5)
    tabela.rows[0]._element.getparent().remove(tabela.rows[0]._element)
    p = documento.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Belo Horizonte\n2026")


def classificar_correlacao(rho: float | None) -> str:
    if rho is None:
        return "indefinida"
    absoluto = abs(float(rho))
    intensidade = "desprezível" if absoluto < 0.1 else "fraca" if absoluto < 0.3 else "moderada" if absoluto < 0.5 else "forte"
    direcao = "positiva" if rho > 0 else "negativa" if rho < 0 else "sem direção"
    return f"{direcao} {intensidade}"


def construir_documento(template: Path, saida: Path, df: pd.DataFrame, estatisticas: dict, graficos: dict[str, Path]) -> None:
    documento = Document(template)
    limpar_corpo(documento)
    configurar_estilos(documento)
    adicionar_capa(documento)

    secao = documento.add_section(WD_SECTION.NEW_PAGE)
    secao.header.is_linked_to_previous = False
    secao.footer.is_linked_to_previous = False
    cabecalho = secao.header.paragraphs[0]
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cabecalho.add_run("Laboratório de Experimentação de Software — Lab01S03 — RQ01–RQ10").font.size = Pt(8.5)
    adicionar_numero_pagina(secao.footer.paragraphs[0])

    total = estatisticas["total"]
    rq08 = estatisticas["rq08"]
    rq09 = estatisticas["rq09"]
    rq10 = estatisticas["rq10"]
    corr_principal = rq08["correlacoes_spearman"]["principal_idade_maior_igual_1"]
    corr_sens = rq08["correlacoes_spearman"]["sensibilidade_todas_idades_positivas"]
    amostra8 = rq08["amostra"]
    correlacao9 = rq09["correlacao_spearman"]
    amostra9 = rq09["amostra"]

    documento.add_heading("1. Introdução", level=1)
    adicionar_texto(documento, "O GitHub reúne projetos com diferentes níveis de maturidade, atividade e participação. Este laboratório investiga características dos 1.000 repositórios públicos mais estrelados retornados pela API GraphQL. Estrelas são usadas como definição operacional de popularidade, sem pressupor qualidade técnica.")
    adicionar_texto(documento, "Esta é a versão intermediária da sprint S03. Ela consolida os resultados atualizados das sete questões sugeridas pelo professor e incorpora a RQ08, sobre popularidade e intensidade anual de desenvolvimento, a RQ09, sobre colaboração e publicação de releases, e a RQ10, sobre maturidade e tratamento de issues. A aplicação Streamlit integra mineração, análise, visualização e exportação do CSV.")
    adicionar_texto(documento, "As hipóteses são informais e exploratórias: orientam a leitura descritiva, não foram pré-registradas e não substituem testes confirmatórios.")

    documento.add_heading("1.1 Questões de pesquisa e hipóteses informais", level=2)
    hipoteses = [
        ["RQ01", "Sistemas populares são maduros/antigos?", "H01: a idade mediana é superior a cinco anos."],
        ["RQ02", "Sistemas populares recebem muita contribuição externa?", "H02: a mediana supera 500 PRs mescladas e a distribuição é assimétrica."],
        ["RQ03", "Sistemas populares lançam releases com frequência?", "H03: a maioria possui releases e a mediana acumulada supera 20."],
        ["RQ04", "Sistemas populares são atualizados com frequência?", "H04: a maioria recebeu push nos últimos 30 dias."],
        ["RQ05", "Sistemas populares usam linguagens populares?", "H05: mais da metade usa uma linguagem do Top 10 Octoverse 2025."],
        ["RQ06", "Sistemas populares possuem alto percentual de issues fechadas?", "H06: a mediana de fechamento supera 80% entre projetos com issues."],
        ["RQ07", "Linguagens populares recebem mais contribuição, releases e atualização?", "H07: o grupo Top 10 tem maiores medianas de PRs e releases e menor recência."],
        ["RQ08", "Popularidade está associada à intensidade anual de desenvolvimento?", "H08: estrelas apresentam associação positiva ao menos moderada com PRs/ano e releases/ano."],
        ["RQ09", "Maior intensidade de PRs está associada a maior intensidade de releases?", "H09: espera-se associação positiva entre PRs mescladas/ano e releases/ano."],
        ["RQ10", "Repositórios mais antigos apresentam maior proporção de issues fechadas?", "H10: repositórios mais antigos apresentam maior percentual mediano de issues fechadas, por possuírem processos de manutenção mais consolidados."],
    ]
    adicionar_tabela(documento, ["RQ", "Pergunta", "Hipótese informal"], hipoteses, fonte=8.1)

    documento.add_heading("2. Desenvolvimento", level=1)
    documento.add_heading("2.1 Evolução do trabalho em S01, S02 e S03", level=2)
    adicionar_texto(documento, "Na S01 foram implementadas a consulta GraphQL inicial, a transformação tabular, as RQ01–RQ07 e a primeira interface Streamlit. Na S02, o coletor em lote passou a percorrer cursores até reunir exatamente 1.000 nomes únicos, exportar CSV e registrar o primeiro snapshot do Project.")
    adicionar_texto(documento, "Na S03, a Issue #44 incorporou a RQ08 com taxas anualizadas, correlações, quartis e visualizações. A Issue #45 acrescentou a RQ09, relacionando PRs mescladas/ano e releases/ano. A Issue #46 incorporou a RQ10 sobre idade e percentual de issues fechadas, e a Issue #47 evoluiu a exibição e a busca no Streamlit. O snapshot final do board permanece pendente até a conclusão do trabalho.")
    adicionar_tabela(documento, ["Sprint", "Atividade", "Responsável", "Rastreabilidade"], [
        ["S01", "Coleta inicial, RQ01–RQ07 e Streamlit", "Pedro e Diogo", "Issues e PRs da S01"],
        ["S02", "Paginação para 1.000, CSV, validação e hipóteses", "Pedro e Diogo", "Issues #8–#15 e #22–#38"],
        ["S03", "RQ08 — popularidade e intensidade", "Pedro Henrique Maia Alves", "Issue #44"],
        ["S03", "RQ09 — colaboração e releases", "Lorran Pedro Avelar Xavier", "Issue #45"],
        ["S03", "RQ10 — maturidade e issues", "Diogo", "Issue #46 — concluída"],
        ["S03", "Evolução de exibição e busca", "Diogo", "Issue #47 — concluída"],
    ], fonte=8.4)

    documento.add_heading("2.2 Arquitetura e fluxo implementado", level=2)
    adicionar_lista(documento, [
        "main.py pagina a consulta GraphQL, valida avanço dos cursores e encerra apenas com 1.000 repositórios únicos;",
        "utils/dataframe.py converte o snapshot bruto e calcula métricas derivadas reproduzíveis;",
        "analysis contém módulos independentes de RQ01 a RQ10 e exporta dez JSONs estritos;",
        "app.py apresenta abas para RQ01–RQ10, busca por repositório e download do CSV com as taxas e correlações anualizadas;",
        "gerar_relatorio_docx.py recalcula todos os números exibidos a partir do mesmo CSV canônico.",
    ])

    documento.add_heading("2.3 Inovação: aplicação Streamlit", level=2)
    adicionar_texto(documento, "A inovação do grupo é uma aplicação interativa em Streamlit. A versão atual apresenta a RQ08 com coeficientes, quartis, sensibilidade e gráficos de dispersão, a RQ09 com correlação, quartis de PRs, tendência e identificação de outliers, além da RQ10 com a relação entre idade e fechamento de issues. Os tooltips preservam os valores originais e as taxas anualizadas usam symlog, mantendo observações iguais a zero.")
    adicionar_texto(documento, "A evolução da Issue #47 acrescentou busca e filtragem por repositório às visualizações. A mineração em lote permanece centralizada no main.py, enquanto o Streamlit atua como camada de exploração e comunicação dos resultados.")

    documento.add_heading("3. Metodologia", level=1)
    documento.add_heading("3.1 Seleção e paginação", level=2)
    adicionar_texto(documento, "A população operacional corresponde aos repositórios públicos encontrados por stars:>0 is:public sort:stars-desc. A conexão search é percorrida em lotes de até 50. Cada nova chamada recebe em after o endCursor anterior; hasNextPage, cursor ausente ou repetido e página vazia são verificados. Duplicatas por nameWithOwner são descartadas até atingir exatamente 1.000 projetos.")

    documento.add_heading("3.2 Transformação e métricas", level=2)
    adicionar_tabela(documento, ["RQ", "Métrica operacional"], [
        ["RQ01", "(data da transformação − createdAt) / 365,25, em anos"],
        ["RQ02", "pullRequests(states: MERGED).totalCount; não identifica autoria externa"],
        ["RQ03", "releases.totalCount acumulado"],
        ["RQ04", "dias entre a transformação e pushedAt"],
        ["RQ05", "primaryLanguage comparada ao Top 10 Octoverse 2025"],
        ["RQ06", "100 × issues fechadas / (abertas + fechadas)"],
        ["RQ07", "medianas de PRs, releases e recência por grupo de linguagem"],
        ["RQ08", "PRs/ano e releases/ano, divididos pela idade registrada no CSV"],
        ["RQ09", "Spearman entre PRs mescladas/ano e releases/ano; medianas por quartil de PRs"],
        ["RQ10", "idade × percentual de issues fechadas; Spearman e mediana/Q1/Q3/IQR por faixa etária"],
    ], fonte=8.4)
    adicionar_texto(documento, "A idade é arredondada para duas casas antes da anualização; as duas taxas são arredondadas para quatro casas. Idades não positivas produzem valor nulo, sem resultados numéricos inválidos.")

    documento.add_heading("3.3 Procedimento analítico da RQ08", level=2)
    adicionar_texto(documento, f"A análise principal inclui idade ≥ 1 ano (n = {amostra8['analise_principal']}); a sensibilidade inclui toda idade positiva (n = {amostra8['sensibilidade_idade_positiva']}). Spearman é calculado como Pearson entre postos médios, sem SciPy. Os quartis são formados após ordenação estável por estrelas e nome. Outliers seguem 1,5 × IQR e permanecem em todas as análises.")
    adicionar_texto(documento, "A anualização reduz o efeito mecânico do tempo de existência, mas não controla domínio, tamanho da equipe, governança, automação ou trajetória histórica. As associações não permitem inferência causal.")

    documento.add_heading("3.4 Procedimento analítico da RQ09", level=2)
    adicionar_texto(documento, f"A RQ09 usa os {amostra9['analise_principal']} repositórios com idade ≥ 1 ano. PRs mescladas e releases acumuladas são divididas pela idade registrada no CSV, com quatro casas decimais. Spearman é calculado entre as duas taxas usando postos médios; os quartis são formados após ordenação estável por PRs/ano e nome.")
    adicionar_texto(documento, "Outliers são identificados separadamente nas duas taxas pelo critério de 1,5 × IQR e mantidos na correlação e nos quartis. A linha de tendência é uma regressão linear descritiva sobre as taxas anualizadas e não substitui a associação monotônica de Spearman.")

    documento.add_heading("3.5 Procedimento analítico da RQ10", level=2)
    adicionar_texto(documento, f"A RQ10 utiliza a idade registrada no CSV e o percentual acumulado de issues fechadas. Dos {formatar_numero(total)} projetos, {formatar_numero(rq10['repositorios_com_issues'])} possuem ao menos uma issue e formam a amostra válida; os {formatar_numero(rq10['repositorios_sem_issues'])} sem issues têm percentual estruturalmente indefinido e são excluídos da correlação e dos resumos por faixa.")
    adicionar_texto(documento, "A associação monotônica é estimada por Spearman, calculado como a correlação de Pearson entre postos médios para tratar empates. A distribuição é resumida nas faixas 0–2, >2–5, >5–10, >10–15 e >15 anos, usando quantidade, mediana, primeiro quartil, terceiro quartil e IQR. Outliers seguem 1,5 × IQR e são mantidos na análise.")
    adicionar_texto(documento, "O percentual acumulado de fechamento não mede velocidade, complexidade ou qualidade da resolução. Projetos podem usar rastreadores externos, e a análise transversal não permite atribuir diferenças etárias a processos de manutenção mais consolidados.")

    documento.add_heading("4. Resultados", level=1)
    documento.add_heading("4.1 Validação do snapshot", level=2)
    adicionar_texto(documento, f"O snapshot contém {formatar_numero(total)} linhas, {formatar_numero(df['Nome'].nunique())} nomes e {formatar_numero(df['URL'].nunique())} URLs únicas. As estrelas estão em ordem não crescente. O CSV contém as taxas PRs por ano e Releases por ano, e os dez arquivos JSON de RQ01–RQ10 registram total_repositorios = {formatar_numero(total)}.")

    documento.add_heading("4.2 Visualizações e resultados por RQ", level=2)
    idade = estatisticas["idade"]
    documento.add_heading("4.2.1 RQ01 — Sistemas populares são maduros/antigos?", level=3)
    adicionar_texto(documento, f"A idade mediana foi {formatar_numero(idade['mediana'], 2)} anos e a média {formatar_numero(idade['media'], 2)}. O intervalo observado foi de {formatar_numero(idade['minimo'], 2)} a {formatar_numero(idade['maximo'], 2)} anos. As faixas <2, 2–<5, 5–<10 e ≥10 anos reuniram, respectivamente, {', '.join(formatar_numero(v) for v in estatisticas['faixas_idade'])} projetos.")
    adicionar_figura(documento, graficos["rq01"], 2, "Distribuição da idade dos repositórios.")

    prs = estatisticas["prs"]
    documento.add_heading("4.2.2 RQ02 — Sistemas populares recebem muita contribuição externa?", level=3)
    adicionar_texto(documento, f"A mediana foi {formatar_numero(prs['mediana'])} PRs mescladas, a média {formatar_numero(prs['media'], 2)}, Q1 {formatar_numero(prs['q1'], 2)}, Q3 {formatar_numero(prs['q3'], 2)} e o máximo {formatar_numero(prs['maximo'])}. Há {estatisticas['prs_zero']} projetos com zero e {estatisticas['prs_outliers']['superiores']} outliers superiores.")
    adicionar_tabela(documento, ["Repositório", "PRs mescladas"], [[nome, formatar_numero(valor)] for nome, valor in estatisticas["top_prs"]], fonte=9)
    adicionar_texto(documento, "PRs mescladas aproximam volume de contribuição, mas a consulta não permite afirmar que a contribuição veio de pessoas externas ao projeto.")
    adicionar_figura(documento, graficos["rq02"], 3, "Distribuição por faixa de pull requests mescladas.")

    releases = estatisticas["releases"]
    documento.add_heading("4.2.3 RQ03 — Sistemas populares lançam releases com frequência?", level=3)
    adicionar_texto(documento, f"A mediana acumulada foi {formatar_numero(releases['mediana'])}, a média {formatar_numero(releases['media'], 2)}, Q1 {formatar_numero(releases['q1'], 2)}, Q3 {formatar_numero(releases['q3'], 2)} e o máximo {formatar_numero(releases['maximo'])}. {estatisticas['releases_com']} projetos ({formatar_percentual(estatisticas['releases_com'], total)}) possuem ao menos uma release; {estatisticas['releases_zero']} possuem zero. Foram detectados {estatisticas['releases_outliers']['superiores']} outliers superiores.")
    adicionar_tabela(documento, ["Repositório", "Releases"], [[nome, formatar_numero(valor)] for nome, valor in estatisticas["top_releases"]], fonte=9)
    adicionar_figura(documento, graficos["rq03"], 4, "Distribuição por quantidade acumulada de releases.")

    dias = estatisticas["dias"]
    documento.add_heading("4.2.4 RQ04 — Sistemas populares são atualizados com frequência?", level=3)
    adicionar_texto(documento, f"A mediana desde o último push foi {formatar_numero(dias['mediana'], 2)} dias, a média {formatar_numero(dias['media'], 2)} e o máximo {formatar_numero(dias['maximo'], 2)}. {estatisticas['dias_ate_30']} projetos ({formatar_percentual(estatisticas['dias_ate_30'], total)}) estavam em até 30 dias; {estatisticas['faixas_dias'][-1]} estavam há mais de um ano. O critério IQR identificou {estatisticas['dias_outliers']['superiores']} outliers superiores.")
    adicionar_tabela(documento, ["Recência", "Repositórios"], [[rotulo, valor] for rotulo, valor in zip(["Até 1 dia", ">1 a 7 dias", ">7 a 30 dias", ">30 a 365 dias", ">365 dias"], estatisticas["faixas_dias"])], fonte=9)
    adicionar_figura(documento, graficos["rq04"], 5, "Distribuição do tempo desde o último push.")

    documento.add_heading("4.2.5 RQ05 — Sistemas populares usam linguagens populares?", level=3)
    adicionar_texto(documento, f"Foram observadas {estatisticas['linguagens_distintas_definidas']} linguagens definidas. O Top 10 Octoverse aparece em {estatisticas['top10']} projetos ({formatar_percentual(estatisticas['top10'], total)}) e em {formatar_percentual(estatisticas['top10'], estatisticas['definida'])} daqueles com linguagem definida. {estatisticas['indefinida']} projetos não possuem linguagem primária definida.")
    adicionar_tabela(documento, ["Linguagem", "Repositórios"], [[nome, valor] for nome, valor in estatisticas["linguagens_top"]], fonte=9)
    adicionar_figura(documento, graficos["rq05"], 6, "Participação das linguagens do Top 10 GitHub Octoverse 2025.")

    issues = estatisticas["issues"]
    documento.add_heading("4.2.6 RQ06 — Sistemas populares possuem alto percentual de issues fechadas?", level=3)
    adicionar_texto(documento, f"A taxa foi calculável em {issues['n']} projetos; {estatisticas['issues_sem']} não possuíam issues. Entre os válidos, a mediana foi {formatar_numero(issues['mediana'], 2)}%, a média {formatar_numero(issues['media'], 2)}%, Q1 {formatar_numero(issues['q1'], 2)}%, Q3 {formatar_numero(issues['q3'], 2)}%, mínimo {formatar_numero(issues['minimo'], 2)}% e máximo {formatar_numero(issues['maximo'], 2)}%. {estatisticas['issues_ge80']} ({formatar_percentual(estatisticas['issues_ge80'], issues['n'])}) alcançaram pelo menos 80%.")
    adicionar_figura(documento, graficos["rq06"], 7, "Distribuição do percentual de issues fechadas.")

    rq07 = estatisticas["rq07"]
    documento.add_heading("4.2.7 RQ07 — Linguagens populares recebem mais contribuição, releases e atualizações?", level=3)
    adicionar_texto(documento, f"A comparação exclui {estatisticas['indefinida']} projetos sem linguagem definida. As estatísticas abaixo foram recalculadas do CSV para os mesmos grupos do Octoverse.")
    adicionar_tabela(documento, ["Grupo", "N", "Mediana PRs", "Mediana releases", "Mediana dias"], [[indice, int(linha['n']), formatar_numero(linha['prs']), formatar_numero(linha['releases']), formatar_numero(linha['dias'], 2)] for indice, linha in rq07.iterrows()], fonte=8.5)
    adicionar_figura(documento, graficos["rq07"], 8, "Indicadores medianos por grupo de linguagens.")

    documento.add_heading("4.2.8 RQ08 — Popularidade e intensidade anual de desenvolvimento", level=3)
    principal_pr = corr_principal["estrelas_vs_prs_por_ano"]
    principal_rel = corr_principal["estrelas_vs_releases_por_ano"]
    sens_pr = corr_sens["estrelas_vs_prs_por_ano"]
    sens_rel = corr_sens["estrelas_vs_releases_por_ano"]
    adicionar_texto(documento, f"Na amostra principal (idade ≥ 1 ano; n = {amostra8['analise_principal']}), estrelas × PRs/ano apresentou ρ = {formatar_numero(principal_pr['rho'], 4)} e estrelas × releases/ano ρ = {formatar_numero(principal_rel['rho'], 4)}. Há {amostra8['idade_menor_1_ano']} projetos com menos de um ano; nenhum outlier foi removido.")
    adicionar_tabela(documento, ["Quartil de estrelas", "N", "Limites de estrelas", "Mediana PRs/ano", "Mediana releases/ano"], [[q["quartil"], q["repositorios"], f"{formatar_numero(q['estrelas_minimo'])}–{formatar_numero(q['estrelas_maximo'])}", formatar_numero(q["prs_por_ano_mediana"], 4), formatar_numero(q["releases_por_ano_mediana"], 4)] for q in rq08["resumo_por_quartil_estrelas"]], fonte=8.1)
    adicionar_texto(documento, f"Na sensibilidade com todas as idades positivas (n = {amostra8['sensibilidade_idade_positiva']}), os coeficientes foram ρ = {formatar_numero(sens_pr['rho'], 4)} para PRs/ano e ρ = {formatar_numero(sens_rel['rho'], 4)} para releases/ano. O IQR marcou {rq08['outliers_iqr']['prs_por_ano']['quantidade']} extremos em PRs/ano, {rq08['outliers_iqr']['releases_por_ano']['quantidade']} em releases/ano e {rq08['outliers_iqr']['uniao_repositorios_extremos']} na união.")
    adicionar_figura(documento, graficos["rq08"], 9, "Dispersões de estrelas versus PRs/ano e releases/ano.")

    faixas9 = rq09["resumo_por_quartil_prs"]
    medianas_releases9 = [faixa["releases_por_ano_mediana"] for faixa in faixas9]
    crescimento_monotonico9 = all(
        atual <= seguinte
        for atual, seguinte in zip(medianas_releases9, medianas_releases9[1:])
    )
    rho9 = correlacao9["rho"]
    documento.add_heading(
        "4.2.9 RQ09 — Colaboração e publicação de releases", level=3
    )
    adicionar_texto(documento, f"Na amostra principal (idade ≥ 1 ano; n = {correlacao9['n']}), a correlação entre PRs mescladas/ano e releases/ano foi {classificar_correlacao(rho9)} (ρ = {formatar_numero(rho9, 4)}). As medianas de releases/ano {'cresceram em todos os quartis' if crescimento_monotonico9 else 'não cresceram em todos os quartis'}, de {formatar_numero(medianas_releases9[0], 4)} no Q1 para {formatar_numero(medianas_releases9[-1], 4)} no Q4.")
    adicionar_tabela(
        documento,
        ["Faixa de PRs/ano", "N", "Limites de PRs/ano", "Mediana PRs/ano", "Mediana releases/ano"],
        [
            [
                faixa["faixa"],
                faixa["repositorios"],
                f"{formatar_numero(faixa['prs_por_ano_minimo'], 4)}–{formatar_numero(faixa['prs_por_ano_maximo'], 4)}",
                formatar_numero(faixa["prs_por_ano_mediana"], 4),
                formatar_numero(faixa["releases_por_ano_mediana"], 4),
            ]
            for faixa in faixas9
        ],
        fonte=8.5,
    )
    adicionar_texto(documento, f"O critério de 1,5 × IQR identificou {rq09['outliers_iqr']['prs_por_ano']['quantidade']} extremos em PRs/ano, {rq09['outliers_iqr']['releases_por_ano']['quantidade']} em releases/ano e {rq09['outliers_iqr']['uniao_repositorios_extremos']} repositórios na união; {rq09['outliers_iqr']['extremos_em_ambas_metricas']} foram extremos nas duas taxas. Todos foram mantidos.")
    adicionar_figura(documento, graficos["rq09"], 10, "Dispersão de PRs/ano versus releases/ano, com outliers e linha de tendência.")

    faixas10 = rq10["resumo_por_faixa_etaria"]
    medianas10 = [faixa["Mediana_issues_fechadas"] for faixa in faixas10]
    crescimento_monotonico10 = all(
        atual <= seguinte
        for atual, seguinte in zip(medianas10, medianas10[1:])
    )
    rho10 = rq10["correlacao_spearman"]
    titulo_rq10 = documento.add_heading(
        "4.2.10 RQ10 — Maturidade e tratamento de issues", level=3
    )
    titulo_rq10.paragraph_format.page_break_before = True
    adicionar_texto(documento, f"A análise inclui {formatar_numero(rq10['repositorios_com_issues'])} repositórios com issues e exclui os {formatar_numero(rq10['repositorios_sem_issues'])} projetos sem issues. A correlação entre idade e percentual de issues fechadas foi {classificar_correlacao(rho10)} (ρ = {formatar_numero(rho10, 4)}). Entre os casos válidos, a idade mediana foi {formatar_numero(rq10['idade_mediana_repositorios_com_issues'], 2)} anos e o percentual mediano de fechamento foi {formatar_numero(rq10['percentual_issues_fechadas_mediana'], 2)}%.")
    adicionar_texto(documento, f"As medianas por faixa {'cresceram em todas as transições' if crescimento_monotonico10 else 'não cresceram em todas as transições'}, passando de {formatar_numero(medianas10[0], 2)}% nos projetos de até dois anos para {formatar_numero(medianas10[-1], 2)}% naqueles com mais de quinze anos. Os intervalos interquartis, porém, apresentam sobreposição, e a última faixa contém apenas {formatar_numero(faixas10[-1]['Quantidade_repositorios'])} projetos.")
    adicionar_tabela(
        documento,
        ["Faixa de idade", "N", "Mediana", "Q1", "Q3", "IQR"],
        [
            [
                faixa["Faixa de idade"],
                faixa["Quantidade_repositorios"],
                f"{formatar_numero(faixa['Mediana_issues_fechadas'], 2)}%",
                f"{formatar_numero(faixa['Q1_issues_fechadas'], 2)}%",
                f"{formatar_numero(faixa['Q3_issues_fechadas'], 2)}%",
                formatar_numero(faixa["IQR_issues_fechadas"], 2),
            ]
            for faixa in faixas10
        ],
        fonte=8.5,
    )
    adicionar_texto(documento, f"No percentual de fechamento, o critério global de 1,5 × IQR identificou {formatar_numero(estatisticas['issues_outliers']['inferiores'])} outliers inferiores e {formatar_numero(estatisticas['issues_outliers']['superiores'])} superiores; todos foram mantidos. Há ainda {formatar_numero(estatisticas['issues_100'])} projetos no teto de 100% de issues fechadas.")
    adicionar_figura(documento, graficos["rq10"], 11, "Relação entre idade e percentual de issues fechadas, com medianas e IQR por faixa etária.")

    documento.add_heading("4.3 Discussão e avaliação das hipóteses", level=2)
    textos_discussao = [
        ["4.3.1 RQ01", f"A mediana de {formatar_numero(idade['mediana'], 2)} anos {'superou' if idade['mediana'] > 5 else 'não superou'} cinco anos. Projetos antigos tiveram mais tempo para acumular estrelas; o resultado não estabelece causalidade."],
        ["4.3.2 RQ02", f"A mediana de {formatar_numero(prs['mediana'])} {'superou' if prs['mediana'] > 500 else 'não superou'} 500. A diferença entre média e mediana e os {estatisticas['prs_outliers']['superiores']} extremos mostram assimetria. A autoria externa não é identificável."],
        ["4.3.3 RQ03", f"{formatar_percentual(estatisticas['releases_com'], total)} possui release e a mediana é {formatar_numero(releases['mediana'])}. O acumulado não mede intervalos nem frequência temporal."],
        ["4.3.4 RQ04", f"{formatar_percentual(estatisticas['dias_ate_30'], total)} estava em até 30 dias. O último push mede recência, não frequência histórica, e pode refletir automação."],
        ["4.3.5 RQ05", f"O Top 10 representa {formatar_percentual(estatisticas['top10'], total)} da amostra. O resultado depende do ranking temporal e primaryLanguage simplifica projetos multilíngues."],
        ["4.3.6 RQ06", f"A mediana de {formatar_numero(issues['mediana'], 2)}% descreve somente projetos com issues. A taxa não informa tempo, complexidade ou qualidade do fechamento."],
        ["4.3.7 RQ07", "A comparação é descritiva. Diferenças entre grupos podem refletir idade, domínio, comunidade, governança e tamanhos desiguais, não o efeito da linguagem isoladamente."],
        ["4.3.8 RQ08", f"A relação com PRs/ano foi {classificar_correlacao(principal_pr['rho'])} (ρ = {formatar_numero(principal_pr['rho'], 4)}) e com releases/ano foi {classificar_correlacao(principal_rel['rho'])} (ρ = {formatar_numero(principal_rel['rho'], 4)}). Assim, H08 não é sustentada no critério de duas associações positivas ao menos moderadas. A sensibilidade manteve a mesma interpretação geral."],
        ["4.3.9 RQ09", f"A H09 recebeu apoio descritivo: a associação foi {classificar_correlacao(rho9)} (ρ = {formatar_numero(rho9, 4)}) e as medianas de releases/ano {'cresceram' if crescimento_monotonico9 else 'não cresceram'} ao longo dos quartis de PRs/ano. O desenho transversal não demonstra que elevar a colaboração cause mais releases."],
        ["4.3.10 RQ10", f"A H10 recebeu apoio descritivo parcial. As medianas {'cresceram monotonicamente' if crescimento_monotonico10 else 'não cresceram monotonicamente'} de {formatar_numero(medianas10[0], 2)}% para {formatar_numero(medianas10[-1], 2)}%, mas a associação foi apenas {classificar_correlacao(rho10)} (ρ = {formatar_numero(rho10, 4)}) e os intervalos interquartis se sobrepõem. A idade isoladamente explica uma parcela limitada da variação e o desenho não permite inferir causalidade."],
    ]
    for titulo, texto in textos_discussao:
        documento.add_heading(titulo, level=3)
        adicionar_texto(documento, texto)

    top_linha = rq07.loc["Top 10 Octoverse"]
    outra_linha = rq07.loc["Outras definidas"]
    avaliacoes = [
        ["H01", "Compatível" if idade["mediana"] > 5 else "Não compatível"],
        ["H02", "Compatível descritivamente" if prs["mediana"] > 500 and estatisticas["prs_outliers"]["superiores"] else "Não compatível"],
        ["H03", "Compatível para presença/acumulado" if estatisticas["releases_com"] > total / 2 and releases["mediana"] > 20 else "Não compatível"],
        ["H04", "Compatível para recência" if estatisticas["dias_ate_30"] > total / 2 else "Não compatível"],
        ["H05", "Compatível" if estatisticas["top10"] > total / 2 else "Não compatível"],
        ["H06", "Compatível entre projetos com issues" if issues["mediana"] > 80 else "Não compatível"],
        ["H07", "Compatível descritivamente" if top_linha["prs"] > outra_linha["prs"] and top_linha["releases"] > outra_linha["releases"] and top_linha["dias"] < outra_linha["dias"] else "Parcial ou não compatível"],
        ["H08", "Não sustentada: as duas associações não foram ao menos moderadas" if not (principal_pr["rho"] >= 0.3 and principal_rel["rho"] >= 0.3) else "Compatível descritivamente"],
        ["H09", "Compatível descritivamente" if rho9 is not None and rho9 > 0 and crescimento_monotonico9 else "Não sustentada"],
        ["H10", "Parcialmente sustentada: tendência monotônica, mas associação fraca" if crescimento_monotonico10 and rho10 is not None and 0 < rho10 < 0.3 else "Compatível descritivamente" if crescimento_monotonico10 and rho10 is not None and rho10 >= 0.3 else "Não sustentada"],
    ]
    adicionar_tabela(documento, ["Hipótese", "Avaliação exploratória"], avaliacoes, fonte=8.7)

    documento.add_heading("4.3.11 Ameaças à validade", level=3)
    adicionar_texto(documento, "Validade de construto. Estrelas aproximam popularidade; PRs mescladas não garantem contribuição externa; releases acumuladas e último push não medem frequência histórica; linguagem primária não descreve todo o código. O percentual de issues fechadas é acumulado e não representa velocidade, dificuldade ou qualidade da resolução.", "Validade de construto.")
    adicionar_texto(documento, "Anualização. Dividir acumulados pela idade supõe uma taxa média constante e amplifica valores de projetos muito jovens. Como PRs/ano e releases/ano usam a mesma idade como denominador, parte da associação pode refletir esse fator comum. Por isso, a análise principal exige ao menos um ano.", "Anualização.")
    adicionar_texto(documento, f"Ausências e efeito teto. Os {formatar_numero(rq10['repositorios_sem_issues'])} projetos sem issues podem usar rastreadores externos ou ainda não ter demandas registradas e foram excluídos da RQ10. Outros {formatar_numero(estatisticas['issues_100'])} aparecem no teto de 100%, reduzindo a discriminação entre projetos.", "Ausências e efeito teto.")
    adicionar_texto(documento, "Outliers. Projetos com automação ou fluxos incomuns produzem taxas extremas. Eles foram identificados por 1,5 × IQR e mantidos; a correlação por postos reduz, mas não elimina, sua influência interpretativa.", "Outliers.")
    adicionar_texto(documento, "Arredondamento e agrupamento. Idades e percentuais registrados no CSV são arredondados, o que cria empates tratados por postos médios. As faixas etárias têm tamanhos distintos, e seus intervalos interquartis se sobrepõem.", "Arredondamento e agrupamento.")
    adicionar_texto(documento, "Causalidade e confundimento. O estudo observacional não controla domínio, organização, tamanho da equipe, governança, efeitos de coorte ou sobrevivência. Associação entre estrelas, intensidade, idade e fechamento não demonstra causalidade.", "Causalidade e confundimento.")
    adicionar_texto(documento, "Validade externa e temporal. A amostra cobre somente os 1.000 repositórios públicos mais estrelados no instante da coleta. Resultados não se generalizam automaticamente e podem variar entre snapshots.", "Validade externa e temporal.")

    documento.add_heading("5. Conclusão", level=1)
    adicionar_texto(documento, f"A coleta oficial reuniu {formatar_numero(total)} repositórios únicos, ordenados por estrelas, e gerou CSV e dez JSONs coerentes para RQ01–RQ10. Todos os valores do relatório foram recalculados do mesmo snapshot, reduzindo o risco de divergência entre dados, texto, tabelas e gráficos.")
    adicionar_texto(documento, f"RQ01–RQ07 descrevem maturidade, contribuições, releases, atualização, linguagens e issues. A RQ08 acrescenta controle simples por idade: na amostra principal, ρ foi {formatar_numero(principal_pr['rho'], 4)} para PRs/ano e {formatar_numero(principal_rel['rho'], 4)} para releases/ano. Os resultados não sustentam a hipótese de duas associações positivas ao menos moderadas e não autorizam conclusão causal.")
    adicionar_texto(documento, f"A RQ09 respondeu positivamente à pergunta de pesquisa: PRs mescladas/ano e releases/ano apresentaram associação {classificar_correlacao(rho9)} (ρ = {formatar_numero(rho9, 4)}), acompanhada pelo crescimento das medianas de releases/ano entre Q1 e Q4. A evidência é associativa e não causal.")
    adicionar_texto(documento, f"Na RQ10, as medianas de issues fechadas aumentaram de {formatar_numero(medianas10[0], 2)}% para {formatar_numero(medianas10[-1], 2)}% entre as faixas extremas, mas a correlação foi positiva fraca (ρ = {formatar_numero(rho10, 4)}). A H10 recebeu apoio descritivo parcial, sem evidência de que a idade cause maior capacidade de fechamento.")
    adicionar_texto(documento, "O Streamlit comunica as RQ01–RQ10, oferece busca por repositório e download do CSV. Esta versão permanece intermediária somente quanto à evidência final do processo: o anexo com o print do fluxo completo do grupo no GitHub Project/Kanban deve ser preenchido após a conclusão de todo o trabalho.")

    documento.add_heading("6. Referências", level=1)
    referencias = [
        "GITHUB. GraphQL API documentation. https://docs.github.com/en/graphql. Acesso em: 25 ago. 2026.",
        "GITHUB. Using pagination in the GraphQL API. https://docs.github.com/en/graphql/guides/using-pagination-in-the-graphql-api. Acesso em: 25 ago. 2026.",
        "GITHUB. Octoverse 2025. https://github.blog/news-insights/octoverse/. Acesso em: 25 ago. 2026.",
        "KALLIAMVAKOU, E. et al. The promises and perils of mining GitHub. MSR, 2014. https://doi.org/10.1145/2597073.2597074.",
        "PANDAS. Pandas documentation. https://pandas.pydata.org/docs/.",
        "STREAMLIT. Streamlit documentation. https://docs.streamlit.io/.",
        "VEGA-ALTAIR. Declarative visualization in Python. https://altair-viz.github.io/.",
    ]
    for referencia in referencias:
        adicionar_texto(documento, referencia)

    titulo_anexo = documento.add_heading("Anexo A — Fluxo do GitHub Project", level=1)
    titulo_anexo.paragraph_format.page_break_before = True
    adicionar_texto(documento, "PENDÊNCIA DA ISSUE #48: capturar somente ao final do trabalho o GitHub Project/Kanban mostrando o fluxo completo do grupo. O estado atual não substitui a captura final.")
    tabela = documento.add_table(rows=1, cols=1)
    tabela.style = "Table Grid"
    celula = tabela.cell(0, 0)
    celula.height = Cm(7)
    celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    sombrear(celula, "F5F7FA")
    p = celula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CAPTURAR AO FINAL: SNAPSHOT DO BOARD COM O FLUXO COMPLETO DO GRUPO").bold = True

    propriedades = documento.core_properties
    propriedades.title = "Lab01S03 — Relatório intermediário RQ01–RQ10"
    propriedades.subject = "Mineração de 1.000 repositórios populares do GitHub"
    propriedades.author = "Pedro Henrique Maia Alves; Diogo C. Brunoro; Lorran Pedro Avelar Xavier"
    propriedades.keywords = "GitHub, GraphQL, Streamlit, RQ08, RQ09, RQ10, releases, issues, Spearman, Lab01S03"
    saida.parent.mkdir(parents=True, exist_ok=True)
    documento.save(saida)


def validar_paridade_rq09_json(rq09: dict, caminho_json: Path) -> None:
    """Garante que o recálculo da RQ09 coincide com o JSON persistido."""

    if not caminho_json.exists():
        raise RuntimeError(f"JSON da RQ09 não encontrado: {caminho_json}")
    esperado = json.loads(caminho_json.read_text(encoding="utf-8"))
    if rq09 != esperado:
        raise RuntimeError("A RQ09 recalculada do CSV diverge do JSON persistido.")


def validar_paridade_rq10_json(rq10: dict, caminho_json: Path) -> None:
    """Garante que o recálculo a partir do CSV coincide com o JSON persistido."""

    if not caminho_json.exists():
        raise RuntimeError(f"JSON da RQ10 não encontrado: {caminho_json}")
    esperado = json.loads(caminho_json.read_text(encoding="utf-8"))
    campos_inteiros = (
        "total_repositorios",
        "repositorios_com_issues",
        "repositorios_sem_issues",
    )
    for campo in campos_inteiros:
        if rq10[campo] != esperado[campo]:
            raise RuntimeError(
                f"Paridade RQ10 divergente em {campo}: "
                f"CSV={rq10[campo]} JSON={esperado[campo]}"
            )
    campos_float = (
        "correlacao_spearman",
        "idade_mediana_repositorios_com_issues",
        "percentual_issues_fechadas_mediana",
    )
    for campo in campos_float:
        obtido = rq10.get(campo)
        alvo = esperado.get(campo)
        if obtido is None or alvo is None:
            if obtido != alvo:
                raise RuntimeError(f"Paridade RQ10 divergente em {campo}: {obtido} vs {alvo}")
            continue
        if not math.isclose(float(obtido), float(alvo), rel_tol=0, abs_tol=1e-4):
            raise RuntimeError(
                f"Paridade RQ10 divergente em {campo}: CSV={obtido} JSON={alvo}"
            )
    campos_faixa = (
        "Quantidade_repositorios",
        "Mediana_issues_fechadas",
        "Q1_issues_fechadas",
        "Q3_issues_fechadas",
        "IQR_issues_fechadas",
    )
    for obtida, alvo in zip(
        rq10["resumo_por_faixa_etaria"],
        esperado["resumo_por_faixa_etaria"],
        strict=True,
    ):
        if obtida["Faixa de idade"] != alvo["Faixa de idade"]:
            raise RuntimeError(
                "Paridade RQ10 divergente nas faixas etárias: "
                f"{obtida['Faixa de idade']} vs {alvo['Faixa de idade']}"
            )
        for campo in campos_faixa:
            valor_obtido = obtida[campo]
            valor_alvo = alvo[campo]
            if isinstance(valor_obtido, int) or isinstance(valor_alvo, int):
                if int(valor_obtido) != int(valor_alvo):
                    raise RuntimeError(
                        f"Paridade RQ10 divergente em {campo} "
                        f"({obtida['Faixa de idade']}): {valor_obtido} vs {valor_alvo}"
                    )
            elif not math.isclose(
                float(valor_obtido), float(valor_alvo), rel_tol=0, abs_tol=1e-4
            ):
                raise RuntimeError(
                    f"Paridade RQ10 divergente em {campo} "
                    f"({obtida['Faixa de idade']}): {valor_obtido} vs {valor_alvo}"
                )


def localizar_navegador_headless() -> Path:
    """Localiza Chrome ou Edge para impressão headless no Windows."""

    candidatos: list[Path] = []
    for variavel in ("PROGRAMFILES", "PROGRAMFILES(X86)"):
        raiz = os.environ.get(variavel)
        if not raiz:
            continue
        base = Path(raiz)
        candidatos.extend(
            [
                base / "Google/Chrome/Application/chrome.exe",
                base / "Microsoft/Edge/Application/msedge.exe",
            ]
        )
    for nome in ("chrome", "msedge"):
        encontrado = shutil.which(nome)
        if encontrado:
            candidatos.append(Path(encontrado))
    for candidato in candidatos:
        if candidato.exists():
            return candidato
    raise RuntimeError(
        "Navegador Chromium não encontrado. Instale Google Chrome ou Microsoft Edge."
    )


def converter_docx_para_pdf(docx: Path, pdf: Path) -> None:
    """Converte DOCX em PDF Carta via HTML temporário (Mammoth + Chrome headless)."""

    try:
        import mammoth
    except ImportError as exc:
        raise RuntimeError(
            "Dependência ausente: instale com 'pip install mammoth'."
        ) from exc

    with docx.open("rb") as arquivo_docx:
        html = mammoth.convert_to_html(arquivo_docx).value
    html = html.replace(
        "<h1>1. Introdução", "<h1 class='nova-pagina'>1. Introdução", 1
    )
    html = html.replace(
        "<h3>4.2.10 RQ10", "<h3 class='nova-pagina'>4.2.10 RQ10", 1
    )
    html = html.replace(
        "<h1>Anexo A", "<h1 class='nova-pagina'>Anexo A", 1
    )

    estilo = """
    @page { size: letter; margin: 2.0cm 2.2cm; }
    body {
        font-family: Calibri, Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.35;
        color: #202020;
    }
    table {
        border-collapse: collapse;
        width: 100%;
        margin: 0.6em 0;
        break-inside: avoid;
        page-break-inside: avoid;
    }
    td, th { border: 1px solid #d0d0d0; padding: 4px 6px; vertical-align: top; }
    img { max-width: 100%; height: auto; }
    h1, h2, h3 { color: #17365D; break-after: avoid; page-break-after: avoid; }
    h3 + p { break-inside: avoid; page-break-inside: avoid; }
    .nova-pagina { break-before: page; page-break-before: always; }
    p { margin: 0.35em 0 0.6em; orphans: 3; widows: 3; }
    """

    documento_html = (
        "<!DOCTYPE html><html lang='pt-BR'><head>"
        "<meta charset='utf-8'>"
        f"<style>{estilo}</style>"
        "</head><body>"
        f"{html}"
        "</body></html>"
    )

    pdf.parent.mkdir(parents=True, exist_ok=True)
    navegador = localizar_navegador_headless()

    with tempfile.TemporaryDirectory(prefix="lab01s03_pdf_") as pasta:
        html_temp = Path(pasta) / "relatorio_temp.html"
        pdf_temp = Path(pasta) / "relatorio_temp.pdf"
        perfil_temp = Path(pasta) / "chrome-profile"
        perfil_temp.mkdir()
        html_temp.write_text(documento_html, encoding="utf-8")
        uri = Path(html_temp).resolve().as_uri()
        comando = [
            str(navegador),
            "--headless=new",
            "--disable-gpu",
            "--disable-extensions",
            "--no-first-run",
            "--no-pdf-header-footer",
            f"--user-data-dir={perfil_temp.resolve()}",
            f"--print-to-pdf={pdf_temp.resolve()}",
            uri,
        ]
        resultado = subprocess.run(
            comando,
            capture_output=True,
            text=True,
            check=False,
        )
        if (
            resultado.returncode != 0
            or not pdf_temp.exists()
            or pdf_temp.stat().st_size == 0
        ):
            detalhes = (resultado.stderr or resultado.stdout or "").strip()
            raise RuntimeError(
                "Falha ao converter DOCX para PDF via navegador headless."
                + (f" Detalhes: {detalhes}" if detalhes else "")
            )
        shutil.copyfile(pdf_temp, pdf)


def validar_documento_gerado(docx: Path) -> None:
    """Executa checklist textual/estrutural básico no DOCX gerado."""

    documento = Document(docx)
    texto = "\n".join(paragrafo.text for paragrafo in documento.paragraphs)
    texto += "\n" + "\n".join(
        celula.text
        for tabela in documento.tables
        for linha in tabela.rows
        for celula in linha.cells
    )
    tabelas = len(documento.tables)
    imagens = len(documento.inline_shapes)

    termos_obrigatorios = (
        "3.4 Procedimento analítico da RQ09",
        "3.5 Procedimento analítico da RQ10",
        "4.2.9 RQ09",
        "4.2.10 RQ10",
        "4.3.9 RQ09",
        "4.3.10 RQ10",
        "Figura 10",
        "Figura 11",
        "H09",
        "H10",
        "Lorran Pedro Avelar Xavier",
        "CAPTURAR AO FINAL: SNAPSHOT DO BOARD",
    )
    faltantes = [termo for termo in termos_obrigatorios if termo not in texto]
    if faltantes:
        raise RuntimeError(f"Documento incompleto; termos ausentes: {', '.join(faltantes)}")
    if "RQ10 pendente" in texto or "RQ10 — pendente" in texto:
        raise RuntimeError("O documento ainda menciona RQ10 como pendente.")
    if "RQ09 pendente" in texto or "RQ09 — pendente" in texto:
        raise RuntimeError("O documento ainda menciona RQ09 como pendente.")
    if tabelas != 14:
        raise RuntimeError(f"Esperadas 14 tabelas, encontradas {tabelas}.")
    if imagens != 10:
        raise RuntimeError(f"Esperadas 10 imagens, encontradas {imagens}.")
    invalido = re.search(r"(?i)\b(?:nan|inf|infinity|infinito)\b", texto)
    if invalido:
        raise RuntimeError(f"Documento contém valor proibido: {invalido.group(0)}")


def criar_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=TEMPLATE_PADRAO)
    parser.add_argument("--csv", type=Path, default=CSV_PADRAO)
    parser.add_argument("--saida", type=Path, default=SAIDA_PADRAO)
    parser.add_argument(
        "--pdf",
        action="store_true",
        help="Converte o DOCX gerado para PDF Carta via Mammoth e Chrome headless.",
    )
    parser.add_argument(
        "--pdf-saida",
        type=Path,
        default=None,
        help="Destino do PDF (padrão: mesmo nome do DOCX com extensão .pdf).",
    )
    parser.add_argument(
        "--json-rq09",
        type=Path,
        default=JSON_RQ09_PADRAO,
        help="JSON da RQ09 usado na validação de paridade.",
    )
    parser.add_argument(
        "--json-rq10",
        type=Path,
        default=JSON_RQ10_PADRAO,
        help="JSON da RQ10 usado na validação de paridade.",
    )
    return parser


def main() -> None:
    argumentos = criar_parser().parse_args()
    if not argumentos.template.exists():
        raise SystemExit(f"Template não encontrado: {argumentos.template}")
    if not argumentos.csv.exists():
        raise SystemExit(f"CSV não encontrado: {argumentos.csv}")
    df = pd.read_csv(argumentos.csv)
    validar_dados(df)
    estatisticas = calcular_estatisticas(df)
    validar_paridade_rq09_json(estatisticas["rq09"], argumentos.json_rq09)
    validar_paridade_rq10_json(estatisticas["rq10"], argumentos.json_rq10)
    with tempfile.TemporaryDirectory(prefix="lab01s03_figuras_") as pasta:
        graficos = preparar_graficos(df, estatisticas, Path(pasta))
        construir_documento(argumentos.template, argumentos.saida, df, estatisticas, graficos)
    validar_documento_gerado(argumentos.saida)
    print(f"Relatório gerado em: {argumentos.saida.resolve()}")
    if argumentos.pdf:
        pdf_saida = argumentos.pdf_saida or argumentos.saida.with_suffix(".pdf")
        converter_docx_para_pdf(argumentos.saida, pdf_saida)
        print(f"PDF gerado em: {pdf_saida.resolve()}")


if __name__ == "__main__":
    main()
